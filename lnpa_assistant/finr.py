import os
import json
import pickle
import numpy as np
import re
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import warnings
warnings.filterwarnings("ignore")

# Базовые библиотеки для обработки
try:
    from sentence_transformers import SentenceTransformer
    from sklearn.metrics.pairwise import cosine_similarity
    from sklearn.feature_extraction.text import TfidfVectorizer
except ImportError as e:
    print(f"❌ Не установлены необходимые библиотеки: {e}")
    print("Установите: pip install sentence-transformers scikit-learn")
    exit(1)

# Библиотеки для работы с DOCX
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx не установлен. Генерация DOCX недоступна.")

# Опциональные библиотеки для расширенного функционала
try:
    import torch
    from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False
    print("⚠️ Transformers не установлен. Расширенная генерация недоступна.")

class DocumentProcessor:
    """Класс для обработки и парсинга документов"""
    
    @staticmethod
    def extract_qa_pairs(text: str) -> List[Dict[str, str]]:
        """Извлечение пар вопрос-ответ из текста"""
        qa_pairs = []
        
        # Разбиваем текст на блоки по вопросам
        question_blocks = re.split(r'\n(?=Вопрос:)', text.strip())
        
        for block in question_blocks:
            block = block.strip()
            if not block or 'Вопрос:' not in block:
                continue
                
            # Извлекаем вопрос
            question_match = re.search(r'Вопрос:\s*(.*?)(?=\nОтвет:|$)', block, re.DOTALL)
            if not question_match:
                continue
            question = question_match.group(1).strip()
            
            # Извлекаем ответ
            answer_match = re.search(r'Ответ:\s*(.*?)(?=\nДокумент:|$)', block, re.DOTALL)
            if not answer_match:
                continue
            answer = answer_match.group(1).strip()
            
            # Извлекаем метаданные
            document_match = re.search(r'Документ:\s*(.*?)(?=\nСсылка:|$)', block, re.DOTALL)
            document = document_match.group(1).strip() if document_match else ""
            
            link_match = re.search(r'Ссылка:\s*(.*?)$', block, re.MULTILINE)
            link = link_match.group(1).strip() if link_match else ""
            
            if question and answer:
                qa_pairs.append({
                    'question': question,
                    'answer': answer,
                    'document': document,
                    'link': link,
                    'full_text': block
                })
        
        return qa_pairs
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Очистка и нормализация текста"""
        # Убираем лишние пробелы и переносы
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n+', '\n', text)
        return text.strip()
    
    @staticmethod
    def extract_document_metadata(text: str) -> Dict[str, str]:
        """Извлечение метаданных документа"""
        metadata = {}
        
        # Поиск номера закона
        law_number = re.search(r'№\s*(\d+-\w+)', text)
        if law_number:
            metadata['law_number'] = law_number.group(1)
        
        # Поиск даты
        date_match = re.search(r'от\s*(\d{1,2}\s+\w+\s+\d{4})', text)
        if date_match:
            metadata['date'] = date_match.group(1)
        
        # Поиск органа
        if 'Республики Беларусь' in text:
            metadata['authority'] = 'Республика Беларусь'
        
        return metadata

class KnowledgeBase:
    """Основной класс базы знаний"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.embedder = SentenceTransformer(model_name)
        self.documents: List[Dict] = []
        self.embeddings: Optional[np.ndarray] = None
        self.tfidf_vectorizer = TfidfVectorizer(max_features=5000, stop_words=None)
        self.tfidf_matrix = None
        self.index_file = "knowledge_base_index.pkl"
        
        print(f"✅ Инициализирована база знаний с моделью: {model_name}")
    
    def add_document_from_file(self, file_path: str) -> int:
        """Добавление документа из файла"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='windows-1251') as f:
                content = f.read()
        
        qa_pairs = DocumentProcessor.extract_qa_pairs(content)
        metadata = DocumentProcessor.extract_document_metadata(content)
        
        added_count = 0
        for qa in qa_pairs:
            qa.update(metadata)
            qa['source_file'] = os.path.basename(file_path)
            qa['added_date'] = datetime.now().isoformat()
            self.documents.append(qa)
            added_count += 1
        
        print(f"📄 Обработан файл {file_path}: добавлено {added_count} записей")
        return added_count
    
    def load_documents_from_folder(self, folder_path: str) -> None:
        """Загрузка всех документов из папки"""
        if not os.path.exists(folder_path):
            raise FileNotFoundError(f"Папка {folder_path} не найдена")
        
        txt_files = [f for f in os.listdir(folder_path) if f.endswith('.txt')]
        if not txt_files:
            raise ValueError(f"В папке {folder_path} нет .txt файлов")
        
        total_added = 0
        for filename in txt_files:
            file_path = os.path.join(folder_path, filename)
            added = self.add_document_from_file(file_path)
            total_added += added
        
        print(f"📚 Всего загружено {total_added} записей из {len(txt_files)} файлов")
        self.build_index()
    
    def build_index(self) -> None:
        """Построение индекса для поиска"""
        if not self.documents:
            raise ValueError("Нет документов для индексации")
        
        print("🔨 Построение индекса...")
        
        # Создаем тексты для эмбеддингов
        texts = []
        for doc in self.documents:
            # Комбинируем вопрос и ответ для лучшего поиска
            combined_text = f"{doc['question']} {doc['answer']}"
            texts.append(combined_text)
        
        # Строим векторные представления
        self.embeddings = self.embedder.encode(texts, convert_to_numpy=True)
        
        # Строим TF-IDF матрицу для дополнительного поиска
        self.tfidf_matrix = self.tfidf_vectorizer.fit_transform(texts)
        
        print(f"✅ Индекс построен для {len(self.documents)} документов")
    
    def save_index(self, filepath: str = None) -> None:
        """Сохранение индекса"""
        filepath = filepath or self.index_file
        data = {
            'documents': self.documents,
            'embeddings': self.embeddings,
            'tfidf_vectorizer': self.tfidf_vectorizer,
            'tfidf_matrix': self.tfidf_matrix
        }
        with open(filepath, 'wb') as f:
            pickle.dump(data, f)
        print(f"💾 Индекс сохранен в {filepath}")
    
    def load_index(self, filepath: str = None) -> None:
        """Загрузка индекса"""
        filepath = filepath or self.index_file
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Файл индекса {filepath} не найден")
        
        with open(filepath, 'rb') as f:
            data = pickle.load(f)
        
        self.documents = data['documents']
        self.embeddings = data['embeddings']
        self.tfidf_vectorizer = data['tfidf_vectorizer']
        self.tfidf_matrix = data['tfidf_matrix']
        
        print(f"📥 Индекс загружен из {filepath}: {len(self.documents)} документов")
    
    def search(self, query: str, top_k: int = 5, method: str = "combined") -> List[Dict]:
        """Поиск по запросу"""
        if not self.documents or self.embeddings is None:
            raise ValueError("База знаний пуста или не проиндексирована")
        
        query_cleaned = DocumentProcessor.clean_text(query)
        
        if method == "semantic":
            return self._semantic_search(query_cleaned, top_k)
        elif method == "tfidf":
            return self._tfidf_search(query_cleaned, top_k)
        else:  # combined
            return self._combined_search(query_cleaned, top_k)
    
    def _semantic_search(self, query: str, top_k: int) -> List[Dict]:
        """Семантический поиск через эмбеддинги"""
        query_embedding = self.embedder.encode([query], convert_to_numpy=True)
        similarities = cosine_similarity(query_embedding, self.embeddings)[0]
        
        top_indices = similarities.argsort()[-top_k:][::-1]
        
        results = []
        for idx in top_indices:
            result = self.documents[idx].copy()
            result['similarity_score'] = float(similarities[idx])
            result['search_method'] = 'semantic'
            results.append(result)
        
        return results
    
    def _tfidf_search(self, query: str, top_k: int) -> List[Dict]:
        """Поиск через TF-IDF"""
        query_tfidf = self.tfidf_vectorizer.transform([query])
        similarities = cosine_similarity(query_tfidf, self.tfidf_matrix)[0]
        
        top_indices = similarities.argsort()[-top_k:][::-1]
        
        results = []
        for idx in top_indices:
            if similarities[idx] > 0:  # Только релевантные результаты
                result = self.documents[idx].copy()
                result['similarity_score'] = float(similarities[idx])
                result['search_method'] = 'tfidf'
                results.append(result)
        
        return results
    
    def _combined_search(self, query: str, top_k: int) -> List[Dict]:
        """Комбинированный поиск"""
        semantic_results = self._semantic_search(query, top_k)
        tfidf_results = self._tfidf_search(query, top_k)
        
        # Комбинируем результаты с весами
        combined_scores = {}
        
        for result in semantic_results:
            doc_id = id(result['full_text'])
            combined_scores[doc_id] = {
                'document': result,
                'semantic_score': result['similarity_score'],
                'tfidf_score': 0.0
            }
        
        for result in tfidf_results:
            doc_id = id(result['full_text'])
            if doc_id in combined_scores:
                combined_scores[doc_id]['tfidf_score'] = result['similarity_score']
            else:
                combined_scores[doc_id] = {
                    'document': result,
                    'semantic_score': 0.0,
                    'tfidf_score': result['similarity_score']
                }
        
        # Вычисляем комбинированный скор
        for doc_id in combined_scores:
            semantic_weight = 0.7
            tfidf_weight = 0.3
            combined_scores[doc_id]['combined_score'] = (
                semantic_weight * combined_scores[doc_id]['semantic_score'] +
                tfidf_weight * combined_scores[doc_id]['tfidf_score']
            )
        
        # Сортируем по комбинированному скору
        sorted_results = sorted(combined_scores.values(), 
                              key=lambda x: x['combined_score'], 
                              reverse=True)
        
        final_results = []
        for item in sorted_results[:top_k]:
            result = item['document'].copy()
            result['similarity_score'] = item['combined_score']
            result['search_method'] = 'combined'
            final_results.append(result)
        
        return final_results

class AdvancedDocumentGenerator:
    """Расширенный класс для генерации документов в формате DOCX"""
    
    def __init__(self):
        self.document = None
        self.templates = self._load_templates()

    def _load_templates(self) -> Dict[str, str]:
        """Загрузка шаблонов документов (текстовые версии)"""
        return {
            'tz': '''ТЕХНИЧЕСКОЕ ЗАДАНИЕ
на {subject}

1. ОБЩИЕ ПОЛОЖЕНИЯ
{general_provisions}

2. ТРЕБОВАНИЯ К ВЫПОЛНЕНИЮ РАБОТ
{requirements}

3. СРОКИ ВЫПОЛНЕНИЯ
{timeline}

4. ОТВЕТСТВЕННОСТЬ СТОРОН
{responsibility}
''',
            'znz': '''ЗАДАНИЕ НА ЗАКУПКУ

Наименование закупки: {procurement_name}
Заказчик: {customer}
Предмет закупки: {subject}

ТЕХНИЧЕСКИЕ ТРЕБОВАНИЯ:
{technical_requirements}

КОММЕРЧЕСКИЕ УСЛОВИЯ:
{commercial_conditions}
''',
            'kp': '''КОНКУРСНОЕ ПРЕДЛОЖЕНИЕ

Участник: {participant}
Предмет конкурса: {subject}

ТЕХНИЧЕСКОЕ ПРЕДЛОЖЕНИЕ:
{technical_proposal}

КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ:
{commercial_proposal}
'''
        }

    def _set_default_font(self):
        """Установка стандартного шрифта"""
        if not DOCX_AVAILABLE:
            return
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Inches(0.14)  # 12pt

    def _add_header(self, data):
        """Добавление шапки 'Утверждаю'"""
        if not DOCX_AVAILABLE:
            return

        # Пустая строка для отступа
        self.document.add_paragraph()

        # Параграф "Утверждаю"
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("Утверждаю\n").bold = True

        # Должность и организация
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"{data.get('approver_position', 'Главный врач')}\n")
        p.add_run(f"{data.get('organization', 'Учреждение')}\n")

        # Подпись и ФИО
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("___________________ ")
        p.add_run(f"{data.get('approver_name', 'И.И. Иванов')}\n")

        # Дата
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f'«____» ___________ {datetime.now().year} г.')

        # Пустая строка
        self.document.add_paragraph()

    def _add_title(self, title):
        """Добавление названия документа"""
        if not DOCX_AVAILABLE:
            return
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(title.upper()).bold = True
        p.add_run().bold = True

        # Пустая строка
        self.document.add_paragraph()

    def _add_footer(self, data):
        """Добавление подписи разработчика"""
        if not DOCX_AVAILABLE:
            return

        # Пустые строки для отступа
        for _ in range(3):
            self.document.add_paragraph()

        # Блок "Разработал"
        p = self.document.add_paragraph()
        p.add_run("Разработал:\n")

        p = self.document.add_paragraph()
        p.add_run(f"{data.get('developer_position', 'Специалист')}\n")

        p = self.document.add_paragraph()
        p.add_run("___________________ ")
        p.add_run(f"{data.get('developer_name', 'П.С. Петров')}")

        # Место и дата
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(f"{data.get('location', 'Минск')} {datetime.now().year}")

    def create_document(self, doc_type: str, data: Dict) -> str:
        """Создание документа по типу и данным"""
        if DOCX_AVAILABLE:
            return self._create_docx_document(doc_type, data)
        else:
            return self._create_text_document(doc_type, data)

    def _create_text_document(self, doc_type: str, context: Dict[str, str]) -> str:
        """Генерация текстового документа"""
        if doc_type not in self.templates:
            raise ValueError(f"Неизвестный тип документа: {doc_type}")
        
        template = self.templates[doc_type]
        
        # Заполняем шаблон контекстом
        try:
            return template.format(**context)
        except KeyError as e:
            raise ValueError(f"Отсутствует обязательный параметр: {e}")

    def _create_docx_document(self, doc_type: str, data: Dict) -> str:
        """Создание документа в формате DOCX"""
        self.document = Document()
        self._set_default_font()

        if doc_type == "ТЗ":
            return self._create_technical_specification(data)
        elif doc_type == "ЗНЗ":
            return self._create_procurement_order(data)
        elif doc_type == "КП":
            return self._create_competitive_proposal(data)
        else:
            raise ValueError("Неизвестный тип документа")

    def _create_technical_specification(self, data):
        """Создание Технического задания в DOCX"""
        # Шапка
        self._add_header(data)

        # Название
        doc_title = data.get('title', 'Техническое задание')
        self._add_title(doc_title)

        # 1. Общие требования
        p = self.document.add_paragraph()
        p.add_run("1. Общие требования\n").bold = True

        requirements = data.get('general_requirements', [])
        for i, req in enumerate(requirements, 1):
            p = self.document.add_paragraph()
            p.add_run(f"1.{i} {req}")

        # 2. Минимальные технические требования
        p = self.document.add_paragraph()
        p.add_run("2. Минимальные технические требования к товару:\n").bold = True

        # Таблица с техническими требованиями
        table_data = data.get('technical_requirements', [])
        if table_data:
            table = self.document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Наименование'
            hdr_cells[1].text = 'Характеристики'
            hdr_cells[2].text = 'Количество'

            # Данные таблицы
            for item in table_data:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get('name', '')
                row_cells[1].text = item.get('characteristics', '')
                row_cells[2].text = item.get('quantity', '')

        # 3. Требования к качеству
        p = self.document.add_paragraph()
        p.add_run("3. Требования к качеству поставляемого товара\n").bold = True

        quality_requirements = data.get('quality_requirements', [])
        for i, req in enumerate(quality_requirements, 1):
            p = self.document.add_paragraph()
            p.add_run(f"3.{i} {req}")

        # 4. Дополнительные требования
        p = self.document.add_paragraph()
        p.add_run("4. Дополнительные требования\n").bold = True

        additional_requirements = data.get('additional_requirements', [])
        for i, req in enumerate(additional_requirements, 1):
            p = self.document.add_paragraph()
            p.add_run(f"4.{i} {req}")

        # 5. Экономическое обоснование
        p = self.document.add_paragraph()
        p.add_run("5. Экономическое обоснование\n").bold = True

        economic_data = data.get('economic_justification', {})
        p = self.document.add_paragraph()
        p.add_run(f"5.1 Источник финансирования: {economic_data.get('funding_source', 'бюджет')}")

        p = self.document.add_paragraph()
        p.add_run(f"5.2 Вид процедуры закупки: {economic_data.get('procurement_type', 'закупка из одного источника')}")

        p = self.document.add_paragraph()
        p.add_run(f"5.3 Условия оплаты: {economic_data.get('payment_terms', 'по факту поставки')}")

        # Подпись
        self._add_footer(data)

        return "Документ создан успешно"

    def _create_procurement_order(self, data):
        """Создание Задания на закупку в DOCX"""
        self._add_header(data)
        self._add_title(data.get('title', 'Задание на закупку'))

        sections = [
            ("1. Общие требования", data.get('general_requirements', [])),
            ("2. Технические характеристики", data.get('technical_specifications', [])),
            ("3. Требования к качеству", data.get('quality_requirements', [])),
            ("4. Условия и сроки поставки", data.get('delivery_terms', [])),
            ("5. Критерии оценки заявок", data.get('evaluation_criteria', [])),
            ("6. Требования к участникам", data.get('participant_requirements', [])),
        ]

        for section_title, section_content in sections:
            p = self.document.add_paragraph()
            p.add_run(f"{section_title}\n").bold = True

            for i, item in enumerate(section_content, 1):
                p = self.document.add_paragraph()
                p.add_run(f"{section_title.split('.')[0]}.{i} {item}")

        # Форма коммерческого предложения
        p = self.document.add_paragraph()
        p.add_run("7. Форма коммерческого предложения\n").bold = True

        table = self.document.add_table(rows=2, cols=8)
        table.style = 'Table Grid'

        # Заголовки таблицы
        headers = ['№ п/п', 'Наименование предмета закупки', 'Кол-во', 'Ед. измерения',
                   'Страна происхождения товара', 'Цена, бел.руб.', 'Сто-сть, бел.руб.', 'Сто-сть, бел.руб. с НДС']

        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        # Пример строки
        example_cells = table.rows[1].cells
        for i in range(len(headers)):
            example_cells[i].text = 'х'

        self._add_footer(data)
        return "Документ создан успешно"

    def _create_competitive_proposal(self, data):
        """Создание Конкурсного предложения в DOCX"""
        self._add_header(data)
        self._add_title(data.get('title', 'Конкурсное предложение'))

        # Аннотация
        p = self.document.add_paragraph()
        p.add_run("Аннотация\n").bold = True
        p = self.document.add_paragraph()
        p.add_run(data.get('annotation', 'Краткое описание предложения объемом не более четырех листов.'))

        sections = [
            ("Коммерческое предложение", data.get('commercial_proposal', [])),
            ("Техническое предложение", data.get('technical_proposal', [])),
            ("Информация о компании", data.get('company_info', [])),
            ("Гарантийные обязательства", data.get('warranty_obligations', [])),
        ]

        for section_title, section_content in sections:
            p = self.document.add_paragraph()
            p.add_run(f"{section_title}\n").bold = True

            for i, item in enumerate(section_content, 1):
                p = self.document.add_paragraph()
                p.add_run(f"{i}. {item}")

        self._add_footer(data)
        return "Документ создан успешно"

    def save_document(self, filename: str) -> str:
        """Сохранение документа"""
        if self.document and DOCX_AVAILABLE:
            self.document.save(filename)
            return f"Документ сохранен как: {filename}"
        else:
            return "Документ не создан или формат DOCX недоступен"

class AIAssistant:
    """Главный класс AI-ассистента"""
    
    def __init__(self):
        self.kb = KnowledgeBase()
        self.generator = AdvancedDocumentGenerator()
        self.llm_pipeline = None
        
        if TRANSFORMERS_AVAILABLE:
            self._init_llm()
    
    def _init_llm(self):
        """Инициализация языковой модели"""
        try:
            model_name = "sberbank-ai/rugpt3small_based_on_gpt2"
            self.llm_pipeline = pipeline(
                "text-generation", 
                model=model_name,
                tokenizer=model_name,
                device=0 if torch.cuda.is_available() else -1
            )
            print("🤖 Языковая модель загружена")
        except Exception as e:
            print(f"⚠️ Не удалось загрузить языковую модель: {e}")
    
    def setup_knowledge_base(self, documents_folder: str) -> None:
        """Настройка базы знаний"""
        try:
            # Пробуем загрузить существующий индекс
            self.kb.load_index()
            print("📥 Загружен существующий индекс")
        except FileNotFoundError:
            # Строим новый индекс
            print("🔨 Строим новый индекс...")
            self.kb.load_documents_from_folder(documents_folder)
            self.kb.save_index()
    
    def ask(self, question: str, use_llm: bool = False) -> Dict[str, any]:
        """Основной метод для вопросов"""
        # Поиск релевантных документов
        results = self.kb.search(question, top_k=3)
        
        if not results:
            return {
                'answer': "Извините, я не нашел релевантной информации по вашему вопросу.",
                'sources': [],
                'confidence': 0.0
            }
        
        # Формируем ответ на основе найденных документов
        best_match = results[0]
        
        if use_llm and self.llm_pipeline:
            answer = self._generate_llm_answer(question, results)
        else:
            answer = best_match['answer']
        
        return {
            'answer': answer,
            'sources': [
                {
                    'document': r['document'],
                    'link': r['link'],
                    'confidence': r['similarity_score']
                } for r in results
            ],
            'confidence': best_match['similarity_score']
        }
    
    def _generate_llm_answer(self, question: str, context_results: List[Dict]) -> str:
        """Генерация ответа с помощью LLM"""
        context = "\n".join([
            f"Вопрос: {r['question']}\nОтвет: {r['answer']}"
            for r in context_results[:2]
        ])
        
        prompt = f"""Ты — опытный юридический консультант, специализирующийся на законодательстве Республики Беларусь. Твоя задача — давать точные, обоснованные ответы на основе нормативных правовых актов, действующих на территории РБ. Используй приведенные фрагменты из базы знаний, чтобы ответить на вопрос пользователя. Ссылайся на документы, избегай домыслов, соблюдай юридическую точность.

Контекст:
{context}

Вопрос: {question}
Ответ:"""
        
        try:
            response = self.llm_pipeline(
                prompt,
                max_length=len(prompt.split()) + 100,
                num_return_sequences=1,
                temperature=0.7,
                do_sample=True
            )
            
            generated_text = response[0]["generated_text"]
            answer = generated_text.split("Ответ:")[-1].strip()
            return answer
            
        except Exception as e:
            print(f"⚠️ Ошибка генерации LLM: {e}")
            return context_results[0]['answer']
    
    def get_statistics(self) -> Dict[str, any]:
        """Статистика базы знаний"""
        return {
            'total_documents': len(self.kb.documents),
            'unique_sources': len(set(doc['source_file'] for doc in self.kb.documents)),
            'has_embeddings': self.kb.embeddings is not None,
            'has_llm': self.llm_pipeline is not None,
            'docx_available': DOCX_AVAILABLE
        }

# Вспомогательные функции для ввода данных
def input_with_default(prompt: str, default: str = "") -> str:
    """Функция для ввода с подсказкой значения по умолчанию"""
    if default:
        user_input = input(f"{prompt} [{default}]: ").strip()
    else:
        user_input = input(f"{prompt}: ").strip()

    return user_input if user_input else default

def input_list(prompt: str, item_name: str = "пункт") -> List[str]:
    """Функция для ввода списка элементов"""
    print(prompt)
    print("(введите каждый элемент с новой строки, пустая строка для завершения)")

    items = []
    i = 1
    while True:
        item = input(f"{item_name} {i}: ").strip()
        if not item:
            break
        items.append(item)
        i += 1

    return items

def input_technical_requirements() -> List[Dict]:
    """Функция для ввода технических требований в виде таблицы"""
    print("\nВвод технических требований:")
    print("(введите данные для каждого товара, пустая строка для завершения)")

    requirements = []
    i = 1
    while True:
        print(f"\nТовар {i}:")
        name = input("Наименование товара: ").strip()
        if not name:
            break

        characteristics = input("Характеристики: ").strip()
        quantity = input("Количество: ").strip()

        requirements.append({
            "name": name,
            "characteristics": characteristics,
            "quantity": quantity
        })
        i += 1

    return requirements

def input_economic_justification() -> Dict:
    """Функция для ввода экономического обоснования"""
    print("\nВвод экономического обоснования:")

    funding_source = input_with_default("Источник финансирования", "областной бюджет")
    procurement_type = input_with_default("Вид процедуры закупки", "закупка из одного источника")
    payment_terms = input_with_default("Условия оплаты", "по факту поставки товара, в течение 5 банковских дней")

    return {
        "funding_source": funding_source,
        "procurement_type": procurement_type,
        "payment_terms": payment_terms
    }

def get_document_data(doc_type: str) -> Dict:
    """Функция для получения данных для конкретного типа документа"""
    print(f"\n=== Ввод данных для документа '{doc_type}' ===\n")

    # Общие данные для всех документов
    data = {}

    data['approver_position'] = input_with_default("Должность утверждающего", "Главный врач")
    data['organization'] = input_with_default("Название организации", "Учреждение")
    data['approver_name'] = input_with_default("ФИО утверждающего", "И.И. Иванов")
    data['developer_position'] = input_with_default("Должность разработчика", "Специалист")
    data['developer_name'] = input_with_default("ФИО разработчика", "П.С. Петров")
    data['location'] = input_with_default("Место составления", "Минск")
    data['title'] = input_with_default("Название документа", f"{get_doc_full_name(doc_type)}")

    # Данные, специфичные для каждого типа документа
    if doc_type == "ТЗ":
        data['general_requirements'] = input_list("\nВведите общие требования:", "требование")
        data['technical_requirements'] = input_technical_requirements()
        data['quality_requirements'] = input_list("\nВведите требования к качеству:", "требование")
        data['additional_requirements'] = input_list("\nВведите дополнительные требования:", "требование")
        data['economic_justification'] = input_economic_justification()

    elif doc_type == "ЗНЗ":
        data['general_requirements'] = input_list("\nВведите общие требования:", "требование")
        data['technical_specifications'] = input_list("\nВведите технические характеристики:", "характеристика")
        data['quality_requirements'] = input_list("\nВведите требования к качеству:", "требование")
        data['delivery_terms'] = input_list("\nВведите условия и сроки поставки:", "условие")
        data['evaluation_criteria'] = input_list("\nВведите критерии оценки заявок:", "критерий")
        data['participant_requirements'] = input_list("\nВведите требования к участникам:", "требование")

    elif doc_type == "КП":
        data['annotation'] = input_with_default("Введите аннотацию", "Краткое описание предложения объемом не более четырех листов.")
        data['commercial_proposal'] = input_list("\nВведите коммерческое предложение:", "пункт")
        data['technical_proposal'] = input_list("\nВведите техническое предложение:", "пункт")
        data['company_info'] = input_list("\nВведите информацию о компании:", "пункт")
        data['warranty_obligations'] = input_list("\nВведите гарантийные обязательства:", "обязательство")

    return data

def get_doc_full_name(doc_type: str) -> str:
    """Получить полное название документа по его типу"""
    doc_names = {
        "ТЗ": "Техническое задание",
        "ЗНЗ": "Задание на закупку",
        "КП": "Конкурсное предложение"
    }
    return doc_names.get(doc_type, "Документ")

# Основные функции меню
def main_menu():
    """Главное меню приложения"""
    assistant = AIAssistant()
    
    while True:
        print("\n" + "="*60)
        print("🤖 AI-АССИСТЕНТ ДЛЯ РАБОТЫ С ЛПА/ЛНПА")
        print("="*60)
        print("1. 📚 Настроить базу знаний")
        print("2. ❓ Задать вопрос")
        print("3. 🔍 Поиск по базе")
        print("4. 📄 Генерация документа")
        print("5. 📊 Статистика")
        print("6. 💾 Управление индексом")
        print("7. ❌ Выход")
        
        choice = input("\nВыберите действие (1-7): ").strip()
        
        if choice == "1":
            setup_kb_menu(assistant)
        elif choice == "2":
            ask_question_menu(assistant)
        elif choice == "3":
            search_menu(assistant)
        elif choice == "4":
            generate_document_menu(assistant)
        elif choice == "5":
            show_statistics(assistant)
        elif choice == "6":
            manage_index_menu(assistant)
        elif choice == "7":
            print("👋 До свидания!")
            break
        else:
            print("❌ Неверный выбор!")

def setup_kb_menu(assistant: AIAssistant):
    """Меню настройки базы знаний"""
    print("\n📚 НАСТРОЙКА БАЗЫ ЗНАНИЙ")
    folder = input("Введите путь к папке с документами [./document]: ").strip() or "./document"
    
    try:
        assistant.setup_knowledge_base(folder)
        print("✅ База знаний настроена успешно!")
    except Exception as e:
        print(f"❌ Ошибка настройки: {e}")

def ask_question_menu(assistant: AIAssistant):
    """Меню для вопросов"""
    if not assistant.kb.documents:
        print("❌ Сначала настройте базу знаний!")
        return
    
    print("\n❓ РЕЖИМ ВОПРОСОВ")
    print("Введите 'назад' для возврата в главное меню")
    
    while True:
        question = input("\n👤 Ваш вопрос: ").strip()
        
        if question.lower() in ['назад', 'back', 'выход']:
            break
        
        if not question:
            continue
        
        use_llm = input("Использовать LLM для генерации ответа? (y/n) [n]: ").strip().lower() == 'y'
        
        try:
            result = assistant.ask(question, use_llm=use_llm)
            
            print(f"\n🤖 Ответ (уверенность: {result['confidence']:.2%}):")
            print(result['answer'])
            
            if result['sources']:
                print(f"\n📚 Источники ({len(result['sources'])}):")
                for i, source in enumerate(result['sources'], 1):
                    print(f"{i}. {source['document']} (релевантность: {source['confidence']:.2%})")
                    if source['link']:
                        print(f"   🔗 {source['link']}")
            
        except Exception as e:
            print(f"❌ Ошибка: {e}")

def search_menu(assistant: AIAssistant):
    """Меню поиска"""
    if not assistant.kb.documents:
        print("❌ Сначала настройте базу знаний!")
        return
    
    print("\n🔍 РЕЖИМ ПОИСКА")
    query = input("Введите поисковый запрос: ").strip()
    
    if not query:
        return
    
    method = input("Метод поиска (semantic/tfidf/combined) [combined]: ").strip() or "combined"
    top_k = input("Количество результатов [5]: ").strip() or "5"
    
    try:
        top_k = int(top_k)
        results = assistant.kb.search(query, top_k=top_k, method=method)
        
        print(f"\n📋 Найдено {len(results)} результатов:")
        for i, result in enumerate(results, 1):
            print(f"\n{i}. Релевантность: {result['similarity_score']:.2%}")
            print(f"❓ Вопрос: {result['question'][:100]}...")
            print(f"💬 Ответ: {result['answer'][:200]}...")
            if result['document']:
                print(f"📄 Документ: {result['document']}")
            
    except ValueError:
        print("❌ Количество результатов должно быть числом!")
    except Exception as e:
        print(f"❌ Ошибка поиска: {e}")

def generate_document_menu(assistant: AIAssistant):
    """Меню генерации документов"""
    print("\n📄 ГЕНЕРАЦИЯ ДОКУМЕНТОВ")
    print("Доступные типы:")
    print("1. ТЗ - Техническое задание")
    print("2. ЗНЗ - Задание на закупку")
    print("3. КП - Конкурсное предложение")
    
    doc_type = input("Выберите тип документа (ТЗ/ЗНЗ/КП): ").strip().upper()
    
    if doc_type not in ['ТЗ', 'ЗНЗ', 'КП']:
        print("❌ Неверный тип документа!")
        return
    
    try:
        # Получаем данные от пользователя
        data = get_document_data(doc_type)
        
        # Выбираем формат
        if DOCX_AVAILABLE:
            format_choice = input("Формат документа (txt/docx) [docx]: ").strip().lower() or "docx"
        else:
            format_choice = "txt"
            print("⚠️ DOCX недоступен, используется текстовый формат")
        
        # Создаем документ
        if format_choice == "docx":
            result = assistant.generator.create_document(doc_type, data)
            filename = f"{doc_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            save_result = assistant.generator.save_document(filename)
            print(f"\n{save_result}")
        else:
            # Текстовый формат
            context = {
                'subject': data.get('title', ''),
                'general_provisions': '\n'.join(data.get('general_requirements', [])),
                'requirements': '\n'.join(data.get('technical_specifications', [])),
                'timeline': data.get('delivery_terms', [''])[0],
                'responsibility': data.get('participant_requirements', [''])[0]
            }
            
            if doc_type == 'ТЗ':
                doc_text = assistant.generator._create_text_document('tz', context)
            elif doc_type == 'ЗНЗ':
                doc_text = assistant.generator._create_text_document('znz', context)
            else:  # КП
                doc_text = assistant.generator._create_text_document('kp', context)
            
            print(f"\n📄 СГЕНЕРИРОВАННЫЙ ДОКУМЕНТ:")
            print("="*60)
            print(doc_text)
            print("="*60)
            
            save = input("\nСохранить документ в файл? (y/n): ").strip().lower()
            if save == 'y':
                filename = input("Имя файла (без расширения): ").strip()
                if filename:
                    with open(f"{filename}.txt", 'w', encoding='utf-8') as f:
                        f.write(doc_text)
                    print(f"✅ Документ сохранен в {filename}.txt")
        
        print(f"\n✓ Документ успешно создан!")
        print(f"Тип: {doc_type} - {get_doc_full_name(doc_type)}")
        print(f"Название: {data.get('title', 'Без названия')}")
        print()
        
    except Exception as e:
        print(f"❌ Ошибка генерации: {e}")
        import traceback
        traceback.print_exc()

def show_statistics(assistant: AIAssistant):
    """Показ статистики"""
    stats = assistant.get_statistics()
    
    print("\n📊 СТАТИСТИКА СИСТЕМЫ")
    print(f"📚 Всего документов: {stats['total_documents']}")
    print(f"📁 Уникальных источников: {stats['unique_sources']}")
    print(f"🔍 Индекс построен: {'✅' if stats['has_embeddings'] else '❌'}")
    print(f"🤖 LLM доступен: {'✅' if stats['has_llm'] else '❌'}")
    print(f"📄 DOCX доступен: {'✅' if stats['docx_available'] else '❌'}")
    
    if assistant.kb.documents:
        # Дополнительная статистика
        sources = {}
        for doc in assistant.kb.documents:
            source = doc['source_file']
            sources[source] = sources.get(source, 0) + 1
        
        print(f"\n📁 Распределение по файлам:")
        for source, count in sources.items():
            print(f"  {source}: {count} записей")

def manage_index_menu(assistant: AIAssistant):
    """Управление индексом"""
    print("\n💾 УПРАВЛЕНИЕ ИНДЕКСОМ")
    print("1. Сохранить индекс")
    print("2. Загрузить индекс")
    print("3. Пересоздать индекс")
    
    choice = input("Выберите действие (1-3): ").strip()
    
    if choice == "1":
        try:
            assistant.kb.save_index()
            print("✅ Индекс сохранен")
        except Exception as e:
            print(f"❌ Ошибка сохранения: {e}")
    
    elif choice == "2":
        try:
            assistant.kb.load_index()
            print("✅ Индекс загружен")
        except Exception as e:
            print(f"❌ Ошибка загрузки: {e}")
    
    elif choice == "3":
        if assistant.kb.documents:
            try:
                assistant.kb.build_index()
                print("✅ Индекс пересоздан")
            except Exception as e:
                print(f"❌ Ошибка пересоздания: {e}")
        else:
            print("❌ Нет документов для индексации")

if __name__ == "__main__":
    print("🚀 Запуск AI-ассистента для работы с ЛПА/ЛНПА...")
    print("📋 Проверка зависимостей...")
    
    # Проверка доступности функций
    if not DOCX_AVAILABLE:
        print("⚠️ Генерация DOCX документов недоступна. Установите: pip install python-docx")
    if not TRANSFORMERS_AVAILABLE:
        print("⚠️ Расширенная генерация ответов недоступна. Установите: pip install transformers torch")
    
    try:
        main_menu()
    except KeyboardInterrupt:
        print("\n\n👋 Работа прервана пользователем")
    except Exception as e:
        print(f"\n❌ Критическая ошибка: {e}")
        import traceback
        traceback.print_exc()

def input_with_default(prompt: str, default: str = "") -> str:
    return default

def input_list(prompt: str, item_name: str = "пункт") -> List[str]:
    return []

def input_technical_requirements() -> List[Dict]:
    return []

def input_economic_justification() -> Dict:
    return {}

def get_document_data(doc_type: str) -> Dict:
    """Функция для получения данных для конкретного типа документа"""
    data = {
        'approver_position': "Главный врач",
        'organization': "Учреждение", 
        'approver_name': "И.И. Иванов",
        'developer_position': "Специалист",
        'developer_name': "П.С. Петров",
        'location': "Минск",
        'title': get_doc_full_name(doc_type)
    }

    if doc_type == "ТЗ":
        data.update({
            'general_requirements': ["Общие требования к системе"],
            'technical_requirements': [{"name": "Пример", "characteristics": "Характеристики", "quantity": "1"}],
            'quality_requirements': ["Требования к качеству"],
            'additional_requirements': ["Дополнительные требования"],
            'economic_justification': {
                'funding_source': "бюджет",
                'procurement_type': "закупка из одного источника", 
                'payment_terms': "по факту поставки"
            }
        })
    elif doc_type == "ЗНЗ":
        data.update({
            'general_requirements': ["Общие требования"],
            'technical_specifications': ["Технические характеристики"],
            'quality_requirements': ["Требования к качеству"],
            'delivery_terms': ["Условия поставки"],
            'evaluation_criteria': ["Критерии оценки"],
            'participant_requirements': ["Требования к участникам"]
        })
    elif doc_type == "КП":
        data.update({
            'annotation': "Краткое описание предложения",
            'commercial_proposal': ["Коммерческое предложение"],
            'technical_proposal': ["Техническое предложение"], 
            'company_info': ["Информация о компании"],
            'warranty_obligations': ["Гарантийные обязательства"]
        })

    return data

def get_doc_full_name(doc_type: str) -> str:
    doc_names = {
        "ТЗ": "Техническое задание",
        "ЗНЗ": "Задание на закупку", 
        "КП": "Конкурсное предложение"
    }
    return doc_names.get(doc_type, "Документ")