import os
import json
import pickle
import numpy as np
import re
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import warnings
warnings.filterwarnings("ignore")

# Базовые библиотеки для обработки
try:
    from sentence_transformers import SentenceTransformer
    from sklearn.metrics.pairwise import cosine_similarity
    from sklearn.feature_extraction.text import TfidfVectorizer
except ImportError as e:
    print(f"❌ Не установлены необходимые библиотеки: {e}")
    print("Установите: pip install sentence-transformers scikit-learn")
    exit(1)

# Библиотеки для работы с DOCX
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx не установлен. Генерация DOCX недоступна.")

# Опциональные библиотеки для расширенного функционала
try:
    import torch
    from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM
    TRANSFORMERS_AVAILABLE = True
except ImportError:
    TRANSFORMERS_AVAILABLE = False
    print("⚠️ Transformers не установлен. Расширенная генерация недоступна.")

class DocumentProcessor:
    """Класс для обработки и парсинга документов"""
    
    @staticmethod
    def extract_qa_pairs(text: str) -> List[Dict[str, str]]:
        """Извлечение пар вопрос-ответ из текста"""
        qa_pairs = []
        
        # Разбиваем текст на блоки по вопросам
        question_blocks = re.split(r'\n(?=Вопрос:)', text.strip())
        
        for block in question_blocks:
            block = block.strip()
            if not block or 'Вопрос:' not in block:
                continue
                
            # Извлекаем вопрос
            question_match = re.search(r'Вопрос:\s*(.*?)(?=\nОтвет:|$)', block, re.DOTALL)
            if not question_match:
                continue
            question = question_match.group(1).strip()
            
            # Извлекаем ответ
            answer_match = re.search(r'Ответ:\s*(.*?)(?=\nДокумент:|$)', block, re.DOTALL)
            if not answer_match:
                continue
            answer = answer_match.group(1).strip()
            
            # Извлекаем метаданные
            document_match = re.search(r'Документ:\s*(.*?)(?=\nСсылка:|$)', block, re.DOTALL)
            document = document_match.group(1).strip() if document_match else ""
            
            link_match = re.search(r'Ссылка:\s*(.*?)$', block, re.MULTILINE)
            link = link_match.group(1).strip() if link_match else ""
            
            if question and answer:
                qa_pairs.append({
                    'question': question,
                    'answer': answer,
                    'document': document,
                    'link': link,
                    'full_text': block
                })
        
        return qa_pairs
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Очистка и нормализация текста"""
        # Убираем лишние пробелы и переносы
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n+', '\n', text)
        return text.strip()
    
    @staticmethod
    def extract_document_metadata(text: str) -> Dict[str, str]:
        """Извлечение метаданных документа"""
        metadata = {}
        
        # Поиск номера закона
        law_number = re.search(r'№\s*(\d+-\w+)', text)
        if law_number:
            metadata['law_number'] = law_number.group(1)
        
        # Поиск даты
        date_match = re.search(r'от\s*(\d{1,2}\s+\w+\s+\d{4})', text)
        if date_match:
            metadata['date'] = date_match.group(1)
        
        # Поиск органа
        if 'Республики Беларусь' in text:
            metadata['authority'] = 'Республика Беларусь'
        
        return metadata

class KnowledgeBase:
    """Основной класс базы знаний"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        self.embedder = SentenceTransformer(model_name)
        self.documents: List[Dict] = []
        self.embeddings: Optional[np.ndarray] = None
        self.tfidf_vectorizer = TfidfVectorizer(max_features=5000, stop_words=None)
        self.tfidf_matrix = None
        self.index_file = "knowledge_base_index.pkl"
        
        print(f"✅ Инициализирована база знаний с моделью: {model_name}")
    
    def add_document_from_file(self, file_path: str) -> int:
        """Добавление документа из файла"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='windows-1251') as f:
                content = f.read()
        
        qa_pairs = DocumentProcessor.extract_qa_pairs(content)
        metadata = DocumentProcessor.extract_document_metadata(content)
        
        added_count = 0
        for qa in qa_pairs:
            qa.update(metadata)
            qa['source_file'] = os.path.basename(file_path)
            qa['added_date'] = datetime.now().isoformat()
            self.documents.append(qa)
            added_count += 1
        
        print(f"📄 Обработан файл {file_path}: добавлено {added_count} записей")
        return added_count
    
    def load_documents_from_folder(self, folder_path: str) -> None:
        """Загрузка всех документов из папки"""
        if not os.path.exists(folder_path):
            raise FileNotFoundError(f"Папка {folder_path} не найдена")
        
        txt_files = [f for f in os.listdir(folder_path) if f.endswith('.txt')]
        if not txt_files:
            raise ValueError(f"В папке {folder_path} нет .txt файлов")
        
        total_added = 0
        for filename in txt_files:
            file_path = os.path.join(folder_path, filename)
            added = self.add_document_from_file(file_path)
            total_added += added
        
        print(f"📚 Всего загружено {total_added} записей из {len(txt_files)} файлов")
        self.build_index()
    
    def build_index(self) -> None:
        """Построение индекса для поиска"""
        if not self.documents:
            raise ValueError("Нет документов для индексации")
        
        print("🔨 Построение индекса...")
        
        # Создаем тексты для эмбеддингов
        texts = []
        for doc in self.documents:
            # Комбинируем вопрос и ответ для лучшего поиска
            combined_text = f"{doc['question']} {doc['answer']}"
            texts.append(combined_text)
        
        # Строим векторные представления
        self.embeddings = self.embedder.encode(texts, convert_to_numpy=True)
        
        # Строим TF-IDF матрицу для дополнительного поиска
        self.tfidf_matrix = self.tfidf_vectorizer.fit_transform(texts)
        
        print(f"✅ Индекс построен для {len(self.documents)} документов")
    
    def save_index(self, filepath: str = None) -> None:
        """Сохранение индекса"""
        filepath = filepath or self.index_file
        data = {
            'documents': self.documents,
            'embeddings': self.embeddings,
            'tfidf_vectorizer': self.tfidf_vectorizer,
            'tfidf_matrix': self.tfidf_matrix
        }
        with open(filepath, 'wb') as f:
            pickle.dump(data, f)
        print(f"💾 Индекс сохранен в {filepath}")
    
    def load_index(self, filepath: str = None) -> None:
        """Загрузка индекса"""
        filepath = filepath or self.index_file
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"Файл индекса {filepath} не найден")
        
        with open(filepath, 'rb') as f:
            data = pickle.load(f)
        
        self.documents = data['documents']
        self.embeddings = data['embeddings']
        self.tfidf_vectorizer = data['tfidf_vectorizer']
        self.tfidf_matrix = data['tfidf_matrix']
        
        print(f"📥 Индекс загружен из {filepath}: {len(self.documents)} документов")
    
    def search(self, query: str, top_k: int = 5, method: str = "combined") -> List[Dict]:
        """Поиск по запросу"""
        if not self.documents or self.embeddings is None:
            raise ValueError("База знаний пуста или не проиндексирована")
        
        query_cleaned = DocumentProcessor.clean_text(query)
        
        if method == "semantic":
            return self._semantic_search(query_cleaned, top_k)
        elif method == "tfidf":
            return self._tfidf_search(query_cleaned, top_k)
        else:  # combined
            return self._combined_search(query_cleaned, top_k)
    
    def _semantic_search(self, query: str, top_k: int) -> List[Dict]:
        """Семантический поиск через эмбеддинги"""
        query_embedding = self.embedder.encode([query], convert_to_numpy=True)
        similarities = cosine_similarity(query_embedding, self.embeddings)[0]
        
        top_indices = similarities.argsort()[-top_k:][::-1]
        
        results = []
        for idx in top_indices:
            result = self.documents[idx].copy()
            result['similarity_score'] = float(similarities[idx])
            result['search_method'] = 'semantic'
            results.append(result)
        
        return results
    
    def _tfidf_search(self, query: str, top_k: int) -> List[Dict]:
        """Поиск через TF-IDF"""
        query_tfidf = self.tfidf_vectorizer.transform([query])
        similarities = cosine_similarity(query_tfidf, self.tfidf_matrix)[0]
        
        top_indices = similarities.argsort()[-top_k:][::-1]
        
        results = []
        for idx in top_indices:
            if similarities[idx] > 0:  # Только релевантные результаты
                result = self.documents[idx].copy()
                result['similarity_score'] = float(similarities[idx])
                result['search_method'] = 'tfidf'
                results.append(result)
        
        return results
    
    def _combined_search(self, query: str, top_k: int) -> List[Dict]:
        """Комбинированный поиск"""
        semantic_results = self._semantic_search(query, top_k)
        tfidf_results = self._tfidf_search(query, top_k)
        
        # Комбинируем результаты с весами
        combined_scores = {}
        
        for result in semantic_results:
            doc_id = id(result['full_text'])
            combined_scores[doc_id] = {
                'document': result,
                'semantic_score': result['similarity_score'],
                'tfidf_score': 0.0
            }
        
        for result in tfidf_results:
            doc_id = id(result['full_text'])
            if doc_id in combined_scores:
                combined_scores[doc_id]['tfidf_score'] = result['similarity_score']
            else:
                combined_scores[doc_id] = {
                    'document': result,
                    'semantic_score': 0.0,
                    'tfidf_score': result['similarity_score']
                }
        
        # Вычисляем комбинированный скор
        for doc_id in combined_scores:
            semantic_weight = 0.7
            tfidf_weight = 0.3
            combined_scores[doc_id]['combined_score'] = (
                semantic_weight * combined_scores[doc_id]['semantic_score'] +
                tfidf_weight * combined_scores[doc_id]['tfidf_score']
            )
        
        # Сортируем по комбинированному скору
        sorted_results = sorted(combined_scores.values(), 
                              key=lambda x: x['combined_score'], 
                              reverse=True)
        
        final_results = []
        for item in sorted_results[:top_k]:
            result = item['document'].copy()
            result['similarity_score'] = item['combined_score']
            result['search_method'] = 'combined'
            final_results.append(result)
        
        return final_results

class AdvancedDocumentGenerator:
    """Расширенный класс для генерации документов в формате DOCX"""
    
    def __init__(self):
        self.document = None
        if not DOCX_AVAILABLE:
            print("⚠️ Генерация DOCX недоступна. Установите python-docx")
    
    def create_document(self, doc_type: str, data: Dict[str, str]) -> 'Document':
        """Создание документа по типу"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx не установлен")
        
        self.document = Document()
        self._set_default_font()

        if doc_type == "ТЗ":
            return self._create_technical_specification(data)
        elif doc_type == "ЗНЗ":
            return self._create_procurement_order(data)
        elif doc_type == "КП":
            return self._create_competitive_proposal(data)
        else:
            raise ValueError("Неизвестный тип документа")

    def _set_default_font(self):
        """Установка стандартного шрифта"""
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Inches(0.14)  # 12pt

    def _add_header(self, data):
        """Добавление шапки 'Утверждаю'"""
        self.document.add_paragraph()

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("Утверждаю\n").bold = True

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"{data.get('approver_position', 'Главный врач')}\n")
        p.add_run(f"{data.get('organization', 'Учреждение')}\n")

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("___________________ ")
        p.add_run(f"{data.get('approver_name', 'И.И. Иванов')}\n")

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f'«____» ___________ {datetime.now().year} г.')

        self.document.add_paragraph()

    def _add_title(self, title):
        """Добавление названия документа"""
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(title.upper()).bold = True
        self.document.add_paragraph()

    def _add_footer(self, data):
        """Добавление подписи разработчика"""
        for _ in range(3):
            self.document.add_paragraph()

        p = self.document.add_paragraph()
        p.add_run("Разработал:\n")

        p = self.document.add_paragraph()
        p.add_run(f"{data.get('developer_position', 'Специалист')}\n")

        p = self.document.add_paragraph()
        p.add_run("___________________ ")
        p.add_run(f"{data.get('developer_name', 'П.С. Петров')}")

        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(f"{data.get('location', 'Минск')} {datetime.now().year}")

    def _create_technical_specification(self, data):
        """Создание Технического задания"""
        self._add_header(data)
        doc_title = data.get('title', 'Техническое задание')
        self._add_title(doc_title)

        p = self.document.add_paragraph()
        p.add_run("1. Общие требования\n").bold = True

        requirements = data.get('general_requirements', [])
        for i, req in enumerate(requirements, 1):
            p = self.document.add_paragraph()
            p.add_run(f"1.{i} {req}")

        p = self.document.add_paragraph()
        p.add_run("2. Минимальные технические требования к товару:\n").bold = True

        table_data = data.get('technical_requirements', [])
        if table_data:
            table = self.document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Наименование'
            hdr_cells[1].text = 'Характеристики'
            hdr_cells[2].text = 'Количество'

            for item in table_data:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get('name', '')
                row_cells[1].text = item.get('characteristics', '')
                row_cells[2].text = item.get('quantity', '')

        p = self.document.add_paragraph()
        p.add_run("3. Требования к качеству поставляемого товара\n").bold = True

        quality_requirements = data.get('quality_requirements', [])
        for i, req in enumerate(quality_requirements, 1):
            p = self.document.add_paragraph()
            p.add_run(f"3.{i} {req}")

        p = self.document.add_paragraph()
        p.add_run("4. Дополнительные требования\n").bold = True

        additional_requirements = data.get('additional_requirements', [])
        for i, req in enumerate(additional_requirements, 1):
            p = self.document.add_paragraph()
            p.add_run(f"4.{i} {req}")

        p = self.document.add_paragraph()
        p.add_run("5. Экономическое обоснование\n").bold = True

        economic_data = data.get('economic_justification', {})
        p = self.document.add_paragraph()
        p.add_run(f"5.1 Источник финансирования: {economic_data.get('funding_source', 'бюджет')}")

        p = self.document.add_paragraph()
        p.add_run(f"5.2 Вид процедуры закупки: {economic_data.get('procurement_type', 'закупка из одного источника')}")

        p = self.document.add_paragraph()
        p.add_run(f"5.3 Условия оплаты: {economic_data.get('payment_terms', 'по факту поставки')}")

        self._add_footer(data)
        return self.document

    def _create_procurement_order(self, data):
        """Создание Задания на закупку"""
        self._add_header(data)
        self._add_title(data.get('title', 'Задание на закупку'))

        sections = [
            ("1. Общие требования", data.get('general_requirements', [])),
            ("2. Технические характеристики", data.get('technical_specifications', [])),
            ("3. Требования к качеству", data.get('quality_requirements', [])),
            ("4. Условия и сроки поставки", data.get('delivery_terms', [])),
            ("5. Критерии оценки заявок", data.get('evaluation_criteria', [])),
            ("6. Требования к участникам", data.get('participant_requirements', [])),
        ]

        for section_title, section_content in sections:
            p = self.document.add_paragraph()
            p.add_run(f"{section_title}\n").bold = True

            for i, item in enumerate(section_content, 1):
                p = self.document.add_paragraph()
                p.add_run(f"{section_title.split('.')[0]}.{i} {item}")

        p = self.document.add_paragraph()
        p.add_run("7. Форма коммерческого предложения\n").bold = True

        table = self.document.add_table(rows=2, cols=8)
        table.style = 'Table Grid'

        headers = ['№ п/п', 'Наименование предмета закупки', 'Кол-во', 'Ед. измерения',
                   'Страна происхождения товара', 'Цена, бел.руб.', 'Сто-сть, бел.руб.', 'Сто-сть, бел.руб. с НДС']

        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header

        example_cells = table.rows[1].cells
        for i in range(len(headers)):
            example_cells[i].text = 'х'

        self._add_footer(data)
        return self.document

    def _create_competitive_proposal(self, data):
        """Создание Конкурсного предложения"""
        self._add_header(data)
        self._add_title(data.get('title', 'Конкурсное предложение'))

        p = self.document.add_paragraph()
        p.add_run("Аннотация\n").bold = True
        p = self.document.add_paragraph()
        p.add_run(data.get('annotation', 'Краткое описание предложения объемом не более четырех листов.'))

        sections = [
            ("Коммерческое предложение", data.get('commercial_proposal', [])),
            ("Техническое предложение", data.get('technical_proposal', [])),
            ("Информация о компании", data.get('company_info', [])),
            ("Гарантийные обязательства", data.get('warranty_obligations', [])),
        ]

        for section_title, section_content in sections:
            p = self.document.add_paragraph()
            p.add_run(f"{section_title}\n").bold = True

            for i, item in enumerate(section_content, 1):
                p = self.document.add_paragraph()
                p.add_run(f"{i}. {item}")

        self._add_footer(data)
        return self.document

    def save_document(self, filename: str) -> str:
        """Сохранение документа"""
        if self.document:
            self.document.save(filename)
            return f"Документ сохранен как: {filename}"
        else:
            return "Документ не создан"

class DocumentGenerator:
    """Класс для генерации документов (расширение системы)"""
    
    def __init__(self, knowledge_base: KnowledgeBase):
        self.kb = knowledge_base
        self.templates = self._load_templates()
        self.advanced_generator = AdvancedDocumentGenerator() if DOCX_AVAILABLE else None
    
    def _load_templates(self) -> Dict[str, str]:
        """Загрузка шаблонов документов"""
        return {
            'tz': '''ТЕХНИЧЕСКОЕ ЗАДАНИЕ
на {subject}

1. ОБЩИЕ ПОЛОЖЕНИЯ
{general_provisions}

2. ТРЕБОВАНИЯ К ВЫПОЛНЕНИЮ РАБОТ
{requirements}

3. СРОКИ ВЫПОЛНЕНИЯ
{timeline}

4. ОТВЕТСТВЕННОСТЬ СТОРОН
{responsibility}
''',
            'znz': '''ЗАДАНИЕ НА ЗАКУПКУ

Наименование закупки: {procurement_name}
Заказчик: {customer}
Предмет закупки: {subject}

ТЕХНИЧЕСКИЕ ТРЕБОВАНИЯ:
{technical_requirements}

КОММЕРЧЕСКИЕ УСЛОВИЯ:
{commercial_conditions}
''',
            'kp': '''КОНКУРСНОЕ ПРЕДЛОЖЕНИЕ

Участник: {participant}
Предмет конкурса: {subject}

ТЕХНИЧЕСКОЕ ПРЕДЛОЖЕНИЕ:
{technical_proposal}

КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ:
{commercial_proposal}
'''
        }
    
    def generate_document(self, doc_type: str, context: Dict[str, str]) -> str:
        """Генерация документа по шаблону"""
        if doc_type not in self.templates:
            raise ValueError(f"Неизвестный тип документа: {doc_type}")
        
        template = self.templates[doc_type]
        
        # Заполняем шаблон контекстом
        try:
            return template.format(**context)
        except KeyError as e:
            raise ValueError(f"Отсутствует обязательный параметр: {e}")
    
    def generate_advanced_document(self, doc_type: str, data: Dict[str, str]) -> str:
        """Генерация расширенного документа в формате DOCX"""
        if not DOCX_AVAILABLE:
            return "Генерация DOCX недоступна. Установите python-docx"
        
        try:
            doc = self.advanced_generator.create_document(doc_type, data)
            filename = f"{doc_type}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            return self.advanced_generator.save_document(filename)
        except Exception as e:
            return f"Ошибка при создании документа: {e}"

class AIAssistant:
    """Главный класс AI-ассистента"""
    
    def __init__(self):
        self.kb = KnowledgeBase()
        self.generator = DocumentGenerator(self.kb)
        self.llm_pipeline = None
        
        if TRANSFORMERS_AVAILABLE:
            self._init_llm()
    
    def _init_llm(self):
        """Инициализация языковой модели"""
        try:
            model_name = "sberbank-ai/rugpt3small_based_on_gpt2"
            self.llm_pipeline = pipeline(
                "text-generation", 
                model=model_name,
                tokenizer=model_name,
                device=0 if torch.cuda.is_available() else -1
            )
            print("🤖 Языковая модель загружена")
        except Exception as e:
            print(f"⚠️ Не удалось загрузить языковую модель: {e}")
    
    def setup_knowledge_base(self, documents_folder: str) -> None:
        """Настройка базы знаний"""
        try:
            # Пробуем загрузить существующий индекс
            self.kb.load_index()
            print("📥 Загружен существующий индекс")
        except FileNotFoundError:
            # Строим новый индекс
            print("🔨 Строим новый индекс...")
            self.kb.load_documents_from_folder(documents_folder)
            self.kb.save_index()
    
    def ask(self, question: str, use_llm: bool = False) -> Dict[str, any]:
        """Основной метод для вопросов"""
        # Поиск релевантных документов
        results = self.kb.search(question, top_k=3)
        
        if not results:
            return {
                'answer': "Извините, я не нашел релевантной информации по вашему вопросу.",
                'sources': [],
                'confidence': 0.0
            }
        
        # Формируем ответ на основе найденных документов
        best_match = results[0]
        
        if use_llm and self.llm_pipeline:
            answer = self._generate_llm_answer(question, results)
        else:
            answer = best_match['answer']
        
        return {
            'answer': answer,
            'sources': [
                {
                    'document': r['document'],
                    'link': r['link'],
                    'confidence': r['similarity_score']
                } for r in results
            ],
            'confidence': best_match['similarity_score']
        }
    
    def _generate_llm_answer(self, question: str, context_results: List[Dict]) -> str:
        """Генерация ответа с помощью LLM"""
        context = "\n".join([
            f"Вопрос: {r['question']}\nОтвет: {r['answer']}"
            for r in context_results[:2]
        ])
        
        prompt = f"""Ты — опытный юридический консультант, специализирующийся на законодательстве Республики Беларусь. Твоя задача — давать точные, обоснованные ответы на основе нормативных правовых актов, действующих на территории РБ. Используй приведенные фрагменты из базы знаний, чтобы ответить на вопрос пользователя. Ссылайся на документы, избегай домыслов, соблюдай юридическую точность.

Контекст:
{context}

Вопрос: {question}
Ответ:"""
        
        try:
            response = self.llm_pipeline(
                prompt,
                max_length=len(prompt.split()) + 100,
                num_return_sequences=1,
                temperature=0.7,
                do_sample=True
            )
            
            generated_text = response[0]["generated_text"]
            answer = generated_text.split("Ответ:")[-1].strip()
            return answer
            
        except Exception as e:
            print(f"⚠️ Ошибка генерации LLM: {e}")
            return context_results[0]['answer']
    
    def get_statistics(self) -> Dict[str, any]:
        """Статистика базы знаний"""
        return {
            'total_documents': len(self.kb.documents),
            'unique_sources': len(set(doc['source_file'] for doc in self.kb.documents)),
            'has_embeddings': self.kb.embeddings is not None,
            'has_llm': self.llm_pipeline is not None,
            'has_docx': DOCX_AVAILABLE
        }

# Вспомогательные функции для ввода данных
def input_with_default(prompt, default=""):
    """Функция для ввода с подсказкой значения по умолчанию"""
    if default:
        user_input = input(f"{prompt} [{default}]: ").strip()
    else:
        user_input = input(f"{prompt}: ").strip()

    return user_input if user_input else default

def input_list(prompt, item_name="пункт"):
    """Функция для ввода списка элементов"""
    print(prompt)
    print("(введите каждый элемент с новой строки, пустая строка для завершения)")

    items = []
    i = 1
    while True:
        item = input(f"{item_name} {i}: ").strip()
        if not item:
            break
        items.append(item)
        i += 1

    return items

def input_technical_requirements():
    """Функция для ввода технических требований в виде таблицы"""
    print("\nВвод технических требований:")
    print("(введите данные для каждого товара, пустая строка для завершения)")

    requirements = []
    i = 1
    while True:
        print(f"\nТовар {i}:")
        name = input("Наименование товара: ").strip()
        if not name:
            break

        characteristics = input("Характеристики: ").strip()
        quantity = input("Количество: ").strip()

        requirements.append({
            "name": name,
            "characteristics": characteristics,
            "quantity": quantity
        })
        i += 1

    return requirements

def input_economic_justification():
    """Функция для ввода экономического обоснования"""
    print("\nВвод экономического обоснования:")

    funding_source = input_with_default("Источник финансирования", "областной бюджет")
    procurement_type = input_with_default("Вид процедуры закупки", "закупка из одного источника")
    payment_terms = input_with_default("Условия оплаты", "по факту поставки товара, в течение 5 банковских дней")

    return {
        "funding_source": funding_source,
        "procurement_type": procurement_type,
        "payment_terms": payment_terms
    }

def get_document_data(doc_type):
    """Функция для получения данных для конкретного типа документа"""
    print(f"\n=== Ввод данных для документа '{doc_type}' ===\n")

    data = {}

    data['approver_position'] = input_with_default("Должность утверждающего", "Главный врач")
    data['organization'] = input_with_default("Название организации", "Учреждение")
    data['approver_name'] = input_with_default("ФИО утверждающего", "И.И. Иванов")
    data['developer_position'] = input_with_default("Должность разработчика", "Специалист")
    data['developer_name'] = input_with_default("ФИО разработчика", "П.С. Петров")
    data['location'] = input_with_default("Место составления", "Минск")
    data['title'] = input_with_default("Название документа", f"{get_doc_full_name(doc_type)}")

    if doc_type == "ТЗ":
        data['general_requirements'] = input_list("\nВведите общие требования:", "требование")
        data['technical_requirements'] = input_technical_requirements()
        data['quality_requirements'] = input_list("\nВведите требования к качеству:", "требование")
        data['additional_requirements'] = input_list("\nВведите дополнительные требования:", "требование")
        data['economic_justification'] = input_economic_justification()

    elif doc_type == "ЗНЗ":
        data['general_requirements'] = input_list("\nВведите общие требования:", "требование")
        data['technical_specifications'] = input_list("\nВведите технические характеристики:", "характеристика")
        data['quality_requirements'] = input_list("\nВведите требования к качеству:", "требование")
        data['delivery_terms'] = input_list("\nВведите условия и сроки поставки:", "условие")
        data['evaluation_criteria'] = input_list("\nВведите критерии оценки заявок:", "критерий")
        data['participant_requirements'] = input_list("\nВведите требования к участникам:", "требование")

    elif doc_type == "КП":
        data['annotation'] = input_with_default("Введите аннотацию", "Краткое описание предложения объемом не более четырех листов.")
        data['commercial_proposal'] = input_list("\nВведите коммерческое предложение:", "пункт")
        data['technical_proposal'] = input_list("\nВведите техническое предложение:", "пункт")
        data['company_info'] = input_list("\nВведите информацию о компании:", "пункт")
        data['warranty_obligations'] = input_list("\nВведите гарантийные обязательства:", "обязательство")

    return data

def get_doc_full_name(doc_type):
    """Получить полное название документа по его типу"""
    doc_names = {
        "ТЗ": "Техническое задание",
        "ЗНЗ": "Задание на закупку",
        "КП": "Конкурсное предложение"
    }
    return doc_names.get(doc_type, "Документ")

def main_menu():
    """Главное меню приложения"""
    assistant = AIAssistant()