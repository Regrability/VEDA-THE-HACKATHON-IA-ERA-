from flask import Flask, request, jsonify, render_template, session, redirect, url_for
from flask_cors import CORS
import os
import sys
import json
import tempfile
from datetime import datetime
import hashlib
import secrets
from functools import wraps

# Добавляем текущую директорию в путь для импорта нашего основного кода
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# Импортируем наш основной код
try:
    from finr import AIAssistant, get_document_data, get_doc_full_name
except ImportError as e:
    print(f"❌ Ошибка импорта finr: {e}")
    # Создаем заглушки для тестирования
    class AIAssistant:
        def __init__(self):
            self.kb = type('obj', (object,), {
                'documents': [
                    {'question': 'Тестовый документ 1', 'answer': 'Тестовый ответ 1', 'source_file': 'test1.txt'},
                    {'question': 'Тестовый документ 2', 'answer': 'Тестовый ответ 2', 'source_file': 'test2.txt'}
                ],
                'add_document_from_file': lambda x: print(f"Добавлен файл: {x}"),
                'build_index': lambda: print("Индекс построен"),
                'search': lambda query, top_k, method: [
                    {'question': 'Результат поиска', 'answer': 'Найденный ответ', 'similarity_score': 0.8, 'source_file': 'search_result.txt'}
                ]
            })()
            self.generator = type('obj', (object,), {
                'create_document': lambda doc_type, data: print(f"Создан документ {doc_type}"),
                'save_document': lambda filename: print(f"Сохранен в {filename}")
            })()
        
        def setup_knowledge_base(self, path):
            print(f"📁 Загрузка базы знаний из {path}")
            
        def ask(self, question, use_llm=False):
            return {
                'answer': f'Ответ на вопрос: {question}',
                'sources': [{'source': 'test.txt', 'page': 1}],
                'confidence': 0.8
            }
        
        def get_statistics(self):
            return {
                'total_documents': len(self.kb.documents),
                'has_llm': True,
                'docx_available': True,
                'has_embeddings': True
            }

# Инициализация Flask приложения
app = Flask(__name__, template_folder='.')
app.secret_key = 'your-secret-key-change-in-production'  # Измените в продакшене
CORS(app)

# Глобальная переменная для ассистента
assistant = None

# Простая база пользователей (в продакшене используйте настоящую БД)
users = {
    'admin': {
        'password': '5e884898da28047151d0e56f8dc6292773603d0d6aabbdd62a11ef721d1542d8',  # password
        'role': 'admin',
        'name': 'Администратор'
    },
    'user': {
        'password': '5e884898da28047151d0e56f8dc6292773603d0d6aabbdd62a11ef721d1542d8',  # password
        'role': 'user',
        'name': 'Обычный пользователь'
    }
}

def hash_password(password):
    """Хеширование пароля"""
    return hashlib.sha256(password.encode()).hexdigest()

def login_required(role='user'):
    """Декоратор для проверки аутентификации и роли"""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if 'user_id' not in session:
                return jsonify({'status': 'error', 'message': 'Требуется авторизация'}), 401
            
            user_role = session.get('role', 'user')
            if role == 'admin' and user_role != 'admin':
                return jsonify({'status': 'error', 'message': 'Недостаточно прав'}), 403
            
            return f(*args, **kwargs)
        return decorated_function
    return decorator

def initialize_assistant():
    """Инициализация AI ассистента"""
    global assistant
    try:
        assistant = AIAssistant()
        # Попытка загрузить существующую базу знаний
        try:
            documents_path = "./documents"
            if not os.path.exists(documents_path):
                os.makedirs(documents_path, exist_ok=True)
                print("📁 Создана папка documents")
            
            assistant.setup_knowledge_base(documents_path)
            print("✅ База знаний загружена успешно")
        except Exception as e:
            print(f"⚠️ Не удалось загрузить базу знаний: {e}")
        
        return True
    except Exception as e:
        print(f"❌ Ошибка инициализации ассистента: {e}")
        return False

@app.route('/')
def index():
    """Главная страница"""
    if 'user_id' in session:
        user_role = session.get('role', 'user')
        return render_template('index.html', user=session, user_role=user_role)
    return redirect(url_for('login_page'))

@app.route('/login')
def login_page():
    """Страница входа"""
    if 'user_id' in session:
        return redirect(url_for('index'))
    return render_template('login.html')

@app.route('/api/login', methods=['POST'])
def login():
    """API для входа"""
    try:
        data = request.get_json()
        username = data.get('username', '').strip()
        password = data.get('password', '').strip()
        
        if not username or not password:
            return jsonify({
                'status': 'error',
                'message': 'Логин и пароль обязательны'
            })
        
        user = users.get(username)
        if user and user['password'] == hash_password(password):
            session['user_id'] = username
            session['role'] = user['role']
            session['name'] = user['name']
            
            return jsonify({
                'status': 'success',
                'data': {
                    'username': username,
                    'role': user['role'],
                    'name': user['name']
                }
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'Неверный логин или пароль'
            })
            
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Ошибка входа: {str(e)}'
        })

@app.route('/api/logout', methods=['POST'])
def logout():
    """API для выхода"""
    session.clear()
    return jsonify({
        'status': 'success',
        'message': 'Выход выполнен'
    })

@app.route('/api/status')
@login_required()
def get_status():
    """Статус системы"""
    if not assistant:
        return jsonify({
            'status': 'error',
            'message': 'Ассистент не инициализирован'
        })
    
    try:
        stats = assistant.get_statistics()
        user_role = session.get('role', 'user')
        return jsonify({
            'status': 'success',
            'data': {
                'assistant_ready': True,
                'documents_count': stats['total_documents'],
                'llm_available': stats['has_llm'],
                'docx_available': stats['docx_available'],
                'index_built': stats['has_embeddings'],
                'user_role': user_role,
                'is_admin': user_role == 'admin'
            }
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        })

@app.route('/api/ask', methods=['POST'])
@login_required()
def ask_question():
    """Обработка вопросов"""
    if not assistant:
        return jsonify({
            'status': 'error',
            'message': 'Ассистент не инициализирован'
        })
    
    try:
        data = request.get_json()
        question = data.get('question', '').strip()
        use_llm = data.get('use_llm', False)
        
        if not question:
            return jsonify({
                'status': 'error',
                'message': 'Вопрос не может быть пустым'
            })
        
        # Получаем ответ от ассистента
        result = assistant.ask(question, use_llm=use_llm)
        
        return jsonify({
            'status': 'success',
            'data': {
                'answer': result['answer'],
                'sources': result['sources'],
                'confidence': result['confidence'],
                'timestamp': datetime.now().isoformat()
            }
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Ошибка обработки вопроса: {str(e)}'
        })

@app.route('/api/search', methods=['POST'])
@login_required()
def search_documents():
    """Поиск по документам"""
    if not assistant:
        return jsonify({
            'status': 'error',
            'message': 'Ассистент не инициализирован'
        })
    
    try:
        data = request.get_json()
        query = data.get('query', '').strip()
        method = data.get('method', 'combined')
        top_k = data.get('top_k', 5)
        
        if not query:
            return jsonify({
                'status': 'error',
                'message': 'Поисковый запрос не может быть пустым'
            })
        
        results = assistant.kb.search(query, top_k=top_k, method=method)
        
        formatted_results = []
        for result in results:
            formatted_results.append({
                'question': result.get('question', ''),
                'answer': result.get('answer', '')[:200] + '...' if len(result.get('answer', '')) > 200 else result.get('answer', ''),
                'document': result.get('document', ''),
                'link': result.get('link', ''),
                'similarity_score': result.get('similarity_score', 0),
                'source_file': result.get('source_file', '')
            })
        
        return jsonify({
            'status': 'success',
            'data': {
                'results': formatted_results,
                'count': len(results),
                'query': query
            }
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Ошибка поиска: {str(e)}'
        })

@app.route('/api/generate-document', methods=['POST'])
@login_required()
def generate_document():
    """Генерация документов"""
    try:
        data = request.get_json()
        doc_type = data.get('doc_type', 'ТЗ')
        document_data = data.get('data', {})
        format_type = data.get('format', 'txt')  # По умолчанию текстовый формат

        print(f"📄 Генерация документа типа: {doc_type}")
        print(f"📋 Полученные данные: {document_data}")
        print(f"📝 Формат: {format_type}")

        # Всегда используем текстовый формат для простоты
        content = generate_text_document(doc_type, document_data)
        
        if format_type == 'docx':
            try:
                # Пытаемся создать DOCX
                from docx import Document
                
                doc = Document()
                doc.add_heading(f'{doc_type}: {document_data.get("title", "Документ")}', 0)
                
                # Разбиваем текстовое содержимое на параграфы
                lines = content.split('\n')
                for line in lines:
                    if line.strip():
                        if line.strip().isupper() and len(line.strip()) > 10:
                            # Заголовок
                            doc.add_heading(line.strip(), level=1)
                        else:
                            # Обычный текст
                            doc.add_paragraph(line.strip())
                
                # Сохраняем во временный файл
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                    filename = tmp_file.name
                    doc.save(filename)
                
                return jsonify({
                    'status': 'success',
                    'data': {
                        'filename': os.path.basename(filename),
                        'format': 'docx',
                        'download_url': f'/api/download/{os.path.basename(filename)}',
                        'message': 'Документ успешно сгенерирован'
                    }
                })
                
            except Exception as e:
                print(f"❌ Ошибка генерации DOCX: {e}")
                # Fallback на текстовый формат
                return jsonify({
                    'status': 'success',
                    'data': {
                        'content': content,
                        'format': 'txt',
                        'message': 'DOCX недоступен, использован текстовый формат'
                    }
                })
        else:
            # Текстовый формат
            return jsonify({
                'status': 'success',
                'data': {
                    'content': content,
                    'format': 'txt'
                }
            })

    except Exception as e:
        print(f"❌ Ошибка генерации документа: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Ошибка генерации: {str(e)}'
        })

def generate_text_document(doc_type, data):
    """Генерация текстового документа"""
    if doc_type == 'ТЗ':
        return generate_tz_document(data)
    elif doc_type == 'ЗНЗ':
        return generate_znz_document(data)
    elif doc_type == 'КП':
        return generate_kp_document(data)
    else:
        return f"Неизвестный тип документа: {doc_type}"

def generate_tz_document(data):
    general_requirements = data.get('general_requirements', [])
    quality_requirements = data.get('quality_requirements', [])
    
    content = f"""ТЕХНИЧЕСКОЕ ЗАДАНИЕ
{data.get('title', 'Техническое задание')}

Утверждаю:
{data.get('approver_position', 'Главный врач')}
{data.get('organization', 'Учреждение')}
___________________ {data.get('approver_name', 'И.И. Иванов')}

1. ОБЩИЕ ТРЕБОВАНИЯ
"""
    
    for i, req in enumerate(general_requirements):
        if req and req.strip():
            content += f"1.{i+1}. {req}\n"
    
    content += "\n2. ТРЕБОВАНИЯ К КАЧЕСТВУ\n"
    for i, req in enumerate(quality_requirements):
        if req and req.strip():
            content += f"2.{i+1}. {req}\n"
    
    content += """
Разработал:
Специалист
___________________ П.С. Петров
"""
    return content

def generate_znz_document(data):
    technical_requirements = data.get('technical_requirements', [])
    
    content = f"""ЗАДАНИЕ НА ЗАКУПКУ
{data.get('title', 'Задание на закупку')}

Заказчик: {data.get('organization', 'Учреждение')}
Предмет закупки: {data.get('subject', 'Медицинское оборудование')}

ТЕХНИЧЕСКИЕ ТРЕБОВАНИЯ:
"""
    
    for req in technical_requirements:
        if req and req.strip():
            content += f"- {req}\n"
    
    return content

def generate_kp_document(data):
    technical_proposal = data.get('technical_proposal', [])
    commercial_proposal = data.get('commercial_proposal', [])
    
    content = f"""КОНКУРСНОЕ ПРЕДЛОЖЕНИЕ
{data.get('title', 'Конкурсное предложение')}

Участник: {data.get('participant', 'ООО Компания')}
Предмет конкурса: {data.get('subject', 'Поставка оборудования')}

ТЕХНИЧЕСКОЕ ПРЕДЛОЖЕНИЕ:
"""
    
    for item in technical_proposal:
        if item and item.strip():
            content += f"- {item}\n"
    
    content += "\nКОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ:\n"
    for item in commercial_proposal:
        if item and item.strip():
            content += f"- {item}\n"
    
    return content

@app.route('/api/download/<filename>')
@login_required()
def download_file(filename):
    """Скачивание сгенерированного файла"""
    try:
        temp_dir = tempfile.gettempdir()
        filepath = os.path.join(temp_dir, filename)
        
        if not os.path.exists(filepath):
            return jsonify({
                'status': 'error',
                'message': 'Файл не найден'
            }), 404
        
        from flask import send_file
        return send_file(filepath, as_attachment=True, download_name=filename)
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Ошибка загрузки файла: {str(e)}'
        }), 500

@app.route('/api/admin/documents', methods=['GET'])
@login_required(role='admin')
def get_documents_list():
    """Получение списка документов (только для админов)"""
    if not assistant:
        return jsonify({
            'status': 'error',
            'message': 'Ассистент не инициализирован'
        })
    
    try:
        documents = getattr(assistant.kb, 'documents', [])
        unique_sources = {}
        
        for doc in documents:
            source = doc.get('source_file', 'unknown')
            if source not in unique_sources:
                unique_sources[source] = {
                    'documents': [],
                    'count': 0
                }
            unique_sources[source]['documents'].append({
                'question': doc.get('question', 'Без названия'),
                'answer_preview': doc.get('answer', '')[:100] + '...' if len(doc.get('answer', '')) > 100 else doc.get('answer', '')
            })
            unique_sources[source]['count'] += 1
        
        return jsonify({
            'status': 'success',
            'data': {
                'total_documents': len(documents),
                'sources': unique_sources
            }
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Ошибка получения списка документов: {str(e)}'
        })

@app.route('/api/admin/documents/<source>', methods=['DELETE'])
@login_required(role='admin')
def delete_document_source(source):
    """Удаление документа из базы знаний (только для админов)"""
    if not assistant:
        return jsonify({
            'status': 'error',
            'message': 'Ассистент не инициализирован'
        })
    
    try:
        # Декодируем имя файла
        import urllib.parse
        source_filename = urllib.parse.unquote(source)
        
        # Удаляем документы из этого источника
        documents = getattr(assistant.kb, 'documents', [])
        initial_count = len(documents)
        assistant.kb.documents = [doc for doc in documents 
                                if doc.get('source_file') != source_filename]
        
        removed_count = initial_count - len(assistant.kb.documents)
        
        # Перестраиваем индекс
        if removed_count > 0 and hasattr(assistant.kb, 'build_index'):
            assistant.kb.build_index()
        
        return jsonify({
            'status': 'success',
            'data': {
                'message': f'Удалено {removed_count} документов из источника {source_filename}',
                'removed_count': removed_count
            }
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Ошибка удаления документов: {str(e)}'
        })

@app.route('/api/upload-documents', methods=['POST'])
@login_required(role='admin')
def upload_documents():
    """Загрузка документов в базу знаний (только для админов)"""
    if not assistant:
        return jsonify({
            'status': 'error',
            'message': 'Ассистент не инициализирован'
        })
    
    try:
        if 'files' not in request.files:
            return jsonify({
                'status': 'error',
                'message': 'Файлы не предоставлены'
            })
        
        files = request.files.getlist('files')
        uploaded_files = []
        skipped_files = []
        
        # Создаем временную папку для загрузки
        upload_dir = os.path.join(tempfile.gettempdir(), 'ai_assistant_uploads')
        os.makedirs(upload_dir, exist_ok=True)
        
        for file in files:
            if file.filename == '':
                continue
            
            # Поддерживаемые форматы
            if file.filename.lower().endswith(('.txt', '.pdf', '.docx', '.doc')):
                filename = os.path.join(upload_dir, file.filename)
                file.save(filename)
                uploaded_files.append(filename)
            else:
                skipped_files.append(file.filename)
        
        # Добавляем документы в базу знаний
        successful_uploads = 0
        for filepath in uploaded_files:
            try:
                if hasattr(assistant.kb, 'add_document_from_file'):
                    assistant.kb.add_document_from_file(filepath)
                    successful_uploads += 1
                    print(f"✅ Успешно загружен: {os.path.basename(filepath)}")
                else:
                    print(f"❌ Метод add_document_from_file не доступен для {filepath}")
            except Exception as e:
                print(f"❌ Ошибка загрузки файла {filepath}: {e}")
        
        # Перестраиваем индекс если есть документы
        if successful_uploads > 0 and hasattr(assistant.kb, 'build_index'):
            assistant.kb.build_index()
            print("✅ Индекс перестроен")
        
        # Удаляем временные файлы
        for filepath in uploaded_files:
            try:
                os.remove(filepath)
            except:
                pass
        
        return jsonify({
            'status': 'success',
            'data': {
                'message': f'Успешно загружено {successful_uploads} документов',
                'uploaded_count': successful_uploads,
                'skipped_files': skipped_files,
                'total_documents': len(assistant.kb.documents) if hasattr(assistant.kb, 'documents') else 0
            }
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Ошибка загрузки документов: {str(e)}'
        })

if __name__ == '__main__':
    print("🚀 Запуск AI Assistant Web Server...")
    
    # Инициализируем ассистента
    if initialize_assistant():
        print("✅ AI Assistant инициализирован успешно")
    else:
        print("❌ Не удалось инициализировать AI Assistant")
    
    # Запускаем сервер
    app.run(debug=True, host='0.0.0.0', port=5000)